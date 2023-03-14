from docx import Document

document = Document()

path = input("What is the file path? ")
readFile = open(path, "r")
Lines = readFile.readlines()

terms = []

def start():
    userInput = input("What section would you like to find? (enter exact name) ")
    searchTerm = userInput.upper()
    terms.append(searchTerm)

    proceed = 0

    while proceed == 0:
        moreYesNo = input("Is there another section you would like to find? (Y/N) ").upper()
        if moreYesNo == "Y":
            extraTerm = input("Please enter a specific section name: ")
            extraTermUpper = extraTerm.upper()
            terms.append(extraTermUpper)
        elif moreYesNo == "N":
            proceed = 1

def findTerms():
    
    for term in terms:
        sections = specifySection(term)
        print(sections)
        lineNo = 0
        termLineNo = []
        lineEmpty =0
        for line in Lines:
            if term in line:
                termLineNo.append(lineNo)     # record the line number where search term is and store it in an array
            lineNo += 1

        for i in sections:
            sectionLines = specifyLines(term, i)
            startLine = termLineNo[i-1]
            lineEmpty = 0

            if sectionLines[0] == 'WHOLE':
                while lineEmpty == 0:
                    if Lines[startLine] != "\n": # TODO: Fix bug on some sections because of empty line condition, likely with function described below near 'LAST'.
                        section = document.add_paragraph(Lines[startLine])
                        startLine += 1
                    else:
                        lineEmpty = 1

            elif sectionLines[0] == 'FIRST':
                lineCount = 0
                while lineCount < int(sectionLines[1]):
                    section = document.add_paragraph(Lines[startLine])
                    startLine += 1
                    lineCount += 1

            elif sectionLines[0] == 'LAST':
                lineCount = 0
                title = document.add_paragraph(term)
                while lineCount < int(sectionLines[1]):
                    section = document.add_paragraph(Lines[startLine+10])  # instead of 10 here, use the return value of a function that gives you how many lines are in a certain section
                    startLine += 1
                    lineCount += 1


def closeAndSave():
    document.save("data_conversion.docx")
    readFile.close()


def specifySection(term):
    sectionInput = input("Which sections of " + term + " would you like to include? (e.g., 1,2) ")
    sections = [int(sectionNumber) for sectionNumber in sectionInput.split(',')] # sections would be [1, 2] if user entered the example above. 
    return sections

def specifyLines(term, sectionNo):
    linesInput = input("Would you like the WHOLE section, FIRST x lines, or LAST x lines of " + term + " #" + str(sectionNo) + "? (e.g., FIRST 5) ").upper()
    sectionLines = linesInput.split(" ")
    # for example, sectionLines would be ['FIRST', '3'] if user input FIRST 3
    return sectionLines


def main():
    start()
    findTerms()
    closeAndSave()
    print(terms) # this print statement is just for debug


if __name__ == "__main__":
    main()
    

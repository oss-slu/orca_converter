from docx import Document
import os


document = Document()
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.font.name = 'Times New Roman'


f = open("/Users/samsam/Desktop/Pyridazine-oxaizridine-m062xdeftzvpJ-optf.txt","r",encoding='UTF-8')
line = f.readline()
while line:
  document.add_paragraph(line)
  line = f.readline()

f.close()

print('Successed！！')

document.save(u'/Users/samsam/Desktop/newdoc.docx')

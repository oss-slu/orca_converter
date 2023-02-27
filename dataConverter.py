from docx import Document
import os

def Convert():
  document = Document()
  paragraph = document.add_paragraph()
  run = paragraph.add_run()
  run.font.name = 'Times New Roman'

  path = input("what is your path?")

  f = open(path,"r",encoding='UTF-8')
  line = f.readline()
  while line:
    line=line.strip()
    document.add_paragraph(line)
    line = f.readline()

  f.close()

  savePath =input("where you gonna save?")
  document.save(savePath)
  print(savePath,'Successed！！')
  return savePath
  
